from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "docs" / "Relatorio_Intermedio_Log_Monitor_MLOps_v2_condense_work.docx"
OUTPUT = ROOT / "docs" / "Relatorio_Intermedio_Log_Monitor_MLOps_v2_condensed_minimal.docx"


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    raise ValueError("Paragraph not found")


def find_paragraphs_exact(doc: Document, text: str) -> list[Paragraph]:
    return [paragraph for paragraph in doc.paragraphs if paragraph_text(paragraph) == text]


def remove_child(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_between(doc: Document, start: Paragraph, end: Paragraph, include_start: bool = False) -> None:
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start._element)
    end_idx = children.index(end._element)
    from_idx = start_idx if include_start else start_idx + 1
    for child in children[from_idx:end_idx]:
        body.remove(child)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        if run.font:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    return new_para


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def insert_table_after(doc: Document, anchor: Paragraph, headers: list[str], rows: list[list[str]]) -> Paragraph:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    anchor._element.addnext(table._tbl)
    spacer = OxmlElement("w:p")
    table._tbl.addnext(spacer)
    return Paragraph(spacer, anchor._parent)


def replace_first_paragraph_starting(doc: Document, startswith: str, new_text: str) -> None:
    paragraph = find_paragraph(doc, lambda p: paragraph_text(p).startswith(startswith))
    paragraph.text = new_text


def remove_table_containing(doc: Document, needle: str) -> None:
    target = None
    for table in doc.tables:
        if any(needle in cell.text for row in table.rows for cell in row.cells):
            target = table
            break
    if target is not None:
        remove_child(target._tbl)


def remove_nth_table_with_header(doc: Document, header_prefix: str, occurrence: int) -> None:
    seen = 0
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip().startswith(header_prefix):
            seen += 1
            if seen == occurrence:
                remove_child(table._tbl)
                return


def trim_table_to_rows_containing(table, keep_row_indices: list[int]) -> None:
    all_rows = list(table.rows)
    for idx, row in enumerate(all_rows):
        if idx not in keep_row_indices:
            remove_child(row._tr)


def find_table_by_header(doc: Document, header_text: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == header_text:
            return table
    raise ValueError(f"Table not found: {header_text}")


def main() -> None:
    doc = Document(INPUT)

    # Front matter: remove template leftovers from list of figures/tables.
    replace_first_paragraph_starting(
        doc,
        "Figura 1 – Processo de carregamento de uma página HTML.",
        "Figura 1 – Diagrama UML — Use Cases Principais.",
    )
    replace_first_paragraph_starting(
        doc,
        "Tabela 1 – Tipos de Selectores existentes.",
        "Tabelas atualizadas no corpo do relatório.",
    )

    # Abstract / summary tone-down.
    replace_first_paragraph_starting(
        doc,
        "The architecture consists of six containerised services",
        "The current architecture runs as a Docker Compose stack with the core ingestion, detection, ML, dashboard, and monitoring services required for local demonstration and validation.",
    )
    replace_first_paragraph_starting(
        doc,
        "At the midpoint of the project (Semanas 1–14), the core system is fully operational.",
        "By Semana 14 of 16, the project already has a demonstrable core covering ingestion, rules, MLflow, a Streamlit dashboard, and monitoring, while some advanced capabilities remain partial or planned.",
    )
    replace_first_paragraph_starting(
        doc,
        "Remaining work for Semanas 15–16 focuses on final documentation",
        "Remaining work for Semanas 15–16 focuses on final documentation, demo polish, and stabilisation. CICIDS-2017 validation, extended feedback loops, and automatic retraining remain planned rather than concluded.",
    )

    # Requirements/use cases: keep the section, but remove the long per-use-case prose.
    use_cases_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "4.3. Casos de Uso Principais")
    actors_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "1.1 Atores do Sistema")
    remove_between(doc, use_cases_heading, actors_heading)
    insert_paragraph_after(
        use_cases_heading,
        "Os casos de uso prioritários nesta fase cobrem ingestão de logs, deteção automática por regras, execução do pipeline híbrido, visualização no dashboard e validação futura em dataset público. O detalhe operacional está resumido na tabela de use cases e no diagrama UML.",
        style="Normal",
    )

    # Viability / state of project.
    replace_first_paragraph_starting(
        doc,
        "As lacunas identificadas no mercado (explicabilidade, feedback automático, custo)",
        "As lacunas identificadas no mercado em explicabilidade, feedback estruturado e custo continuam relevantes. O projeto demonstra viabilidade técnica e académica, mas ainda não deve ser descrito como solução totalmente fechada em feedback ou incident response.",
    )
    replace_first_paragraph_starting(
        doc,
        "Estado atual: O projeto encontra-se na Semana 14 de 16. Todas as funcionalidades do Goal A e Goal B foram concluídas.",
        "Estado atual: O projeto encontra-se na Semana 14 de 16. A maioria das funcionalidades dos Goals A e B foi concluída. Permanecem parciais ou diferidas o feedback operacional completo, o workflow de incidentes, o benchmark externo e o retraining automático.",
    )

    # Solution proposed corrections.
    replace_first_paragraph_starting(
        doc,
        "S1 — Setup do repositório, ambiente Docker, estrutura do projeto: concluído. S2 — Flask app e gerador de tráfego sintético (10 tipos de ataque simulados): concluído.",
        "S1 — Setup do repositório, ambiente Docker e estrutura do projeto: concluído. S2 — Flask app e gerador de tráfego sintético com cenários normais e de ataque principais: concluído. S3 — PostgreSQL + TimescaleDB, pipeline de ingestão com batch insert: concluído. S4 — Rule Engine para deteção determinística de ameaças conhecidas: concluído. S5–S8 — feature engineering, treino ML, SHAP e dashboard inicial: concluído.",
    )
    replace_first_paragraph_starting(
        doc,
        "S9 — Docker Compose com 6 containers",
        "S9 — Docker Compose para a stack principal e integração dos serviços core: concluído. S10 — Stack de monitorização com Prometheus, Grafana e Alertmanager: concluído. S11 — Dashboard Streamlit de operador: concluído. S12 — testing, load test e quality gates: concluído. S13–S14 — CI/CD e security hardening base: concluído.",
    )
    replace_first_paragraph_starting(
        doc,
        "Funcionalidades Diferidas (Pós-Apresentação)",
        "Funcionalidades Diferidas para a Fase Final do Projeto",
    )
    replace_first_paragraph_starting(
        doc,
        "As seguintes funcionalidades foram avaliadas e diferidas por terem ROI baixo face ao esforço no contexto académico atual",
        "As seguintes funcionalidades foram diferidas para a fase final por apresentarem menor ROI na apresentação intercalar ou por dependerem de mais dados e infraestrutura: LIME e counterfactuals; benchmark CICIDS-2017; retraining automático com feedback; cloud deployment; e workflow de incidentes totalmente fechado.",
    )
    replace_first_paragraph_starting(
        doc,
        "Testing & CI (S12–S13): Test coverage: 74%",
        "Testing & CI (S12–S13): Na execução local mais recente do pytest foram recolhidos 102 testes; 87 passaram, 1 ficou skipped, 2 falharam e 12 deram erro, com cobertura global de 70.97%. As falhas concentram-se sobretudo no módulo hybrid_pipeline, pelo que a qualidade do core já é demonstrável, mas ainda não está totalmente estabilizada.",
    )

    # Differentiation / overclaiming corrections.
    replace_first_paragraph_starting(
        doc,
        "Camada 1 - Regras Determinísticas: • Deteção de ataques conhecidos com precisão de 100% (zero FP) • 7 regras configuráveis",
        "Camada 1 - Regras Determinísticas: deteção de ataques conhecidos em cenários sintéticos controlados, com 6 regras implementadas no rule engine (brute force, SQL injection, port scanning, path traversal, suspicious user agent e time-based anomaly).",
    )
    replace_first_paragraph_starting(
        doc,
        "Workflow: 1. Analyst marca alerta como False Positive no dashboard",
        "Workflow planeado: o analyst investiga um alerta, regista feedback estruturado e esse feedback fica disponível para aprendizagem futura. O fluxo operacional completo no dashboard e o retraining automático ainda pertencem à fase final.",
    )
    replace_first_paragraph_starting(
        doc,
        "A recolha de feedback estruturado (FP/FN com razão) está implementada.",
        "A estrutura de feedback está preparada no schema e faz parte da arquitetura do projeto, mas a recolha operacional completa via dashboard e o retraining automático ainda não estão fechados ponta a ponta.",
    )
    replace_first_paragraph_starting(
        doc,
        "O sistema é testado e validado em datasets públicos académicos:",
        "A validação principal nesta fase assenta no dataset interno e em testes de comportamento do sistema. A validação em dataset público académico permanece planeada para a fase final.",
    )
    replace_first_paragraph_starting(
        doc,
        "Isto permite afirmar com confiança: 'O sistema tem performance comparável a métodos publicados em papers académicos'.",
        "Assim, o relatório pode afirmar com segurança que o projeto já possui base experimental interna consistente, mas ainda não deve apresentar benchmark público como resultado concluído.",
    )
    replace_first_paragraph_starting(
        doc,
        "Estados do Incidente: NEW → INVESTIGATING → CONFIRMED → MITIGATING → RESOLVED",
        "Estado atual: o dashboard já suporta investigação e playbooks estáticos. O workflow completo de incidentes com estados, assignment, timeline operacional e KPIs continua parcial e fica reservado para a fase final.",
    )
    replace_first_paragraph_starting(
        doc,
        "Implementado: • Web application logs (HTTP requests) • Docker container logs (CPU, memory, eventos)",
        "Implementado de forma sólida: web application logs (HTTP requests). A arquitetura foi pensada para crescer para outros tipos de logs, mas essa extensão ainda não constitui capacidade operacional fechada nesta fase.",
    )
    # Cut the heaviest low-ROI sections while keeping the document identity.
    benchmark_detail_start = find_paragraph(doc, lambda p: paragraph_text(p) == "2.2. Análise Detalhada por Solução")
    benchmark_gap_candidates = find_paragraphs_exact(doc, "2.3. Gap Analysis")
    benchmark_gap = benchmark_gap_candidates[-1]
    remove_between(doc, benchmark_detail_start, benchmark_gap, include_start=True)

    tech_categories_start = find_paragraph(doc, lambda p: paragraph_text(p) == "3.2 Tecnologias por Categoria")
    tech_justifications = find_paragraph(doc, lambda p: paragraph_text(p) == "3.3 Justificação de Escolhas Críticas")
    remove_between(doc, tech_categories_start, tech_justifications, include_start=True)

    deps_start = find_paragraph(doc, lambda p: paragraph_text(p) == "3.4 Dependências Python Principais")
    benchmarking_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "Benchmarking")
    remove_between(doc, deps_start, benchmarking_heading, include_start=True)

    # Remove low-ROI tables and compress the heaviest planning matrix.
    remove_table_containing(doc, "CICIDS-2017")
    remove_nth_table_with_header(doc, "Aspeto", 2)

    # Replace the large Gantt matrix with a compact Gantt summary.
    gantt_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "2.2 Gantt Detalhado")
    critical_path_heading = find_paragraph(doc, lambda p: paragraph_text(p) == "2.3 Critical Path e Dependências")
    remove_between(doc, gantt_heading, critical_path_heading)
    gantt_intro = insert_paragraph_after(
        gantt_heading,
        "A versão intermédia mantém o Gantt por blocos de semanas, preservando a lógica temporal sem a matriz semanal extensa.",
        style="Normal",
    )
    insert_table_after(
        doc,
        gantt_intro,
        ["Frente", "S1-2", "S3-4", "S5-6", "S7-8", "S9-10", "S11-12", "S13-14", "S15-16"],
        [
            ["Setup e arquitetura base", "X", "", "", "", "", "", "", ""],
            ["Ingestão e persistência", "", "X", "X", "", "", "", "", ""],
            ["Rules + ML core", "", "", "X", "X", "", "", "", ""],
            ["Dashboard + monitoring", "", "", "", "", "X", "X", "", ""],
            ["QA + hardening + CI/CD", "", "", "", "", "", "X", "X", ""],
            ["Documentação e demo final", "", "", "", "", "", "", "", "X"],
        ],
    )

    # Remove actor table and use-case summary table, keeping the surrounding explanatory text.
    remove_table_containing(doc, "Security Analyst")
    remove_table_containing(doc, "UC-01")

    # Trim verbose requirements tables to the most important rows only.
    functional_table = find_table_by_header(doc, "ID")
    trim_table_to_rows_containing(functional_table, [0, 1, 2, 3, 6, 9, 12, 15])
    non_functional_table = doc.tables[1]
    trim_table_to_rows_containing(non_functional_table, [0, 1, 2, 3, 4, 6, 8, 12, 14])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
