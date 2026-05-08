from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "Template_PI_Relatorio_Intermedio_work.docx"
OUTPUT = ROOT / "docs" / "Relatorio_Intermedio_Log_Monitor_MLOps_v5_template_aligned.docx"
MEDIA_DIR = ROOT / "docs" / "report_media"
UML_COMPOSITE = ROOT / "docs" / "report_media" / "uml_use_cases_composite.png"


def combine_images() -> Path:
    paths = [MEDIA_DIR / "image2.png", MEDIA_DIR / "image3.png", MEDIA_DIR / "image4.png"]
    images = [Image.open(path).convert("RGB") for path in paths]
    target_width = 1480
    spacing = 24
    resized = []
    for image in images:
        ratio = target_width / image.width
        size = (target_width, int(image.height * ratio))
        resized.append(image.resize(size))
    total_height = sum(image.height for image in resized) + spacing * (len(resized) - 1)
    canvas = Image.new("RGB", (target_width, total_height), "white")
    y = 0
    for image in resized:
        canvas.paste(image, (0, y))
        y += image.height + spacing
    canvas.save(UML_COMPOSITE)
    return UML_COMPOSITE


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_between(doc: Document, start: Paragraph, end: Paragraph) -> None:
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start._element)
    end_idx = children.index(end._element)
    for child in children[start_idx + 1 : end_idx]:
        body.remove(child)


def remove_after(doc: Document, start: Paragraph) -> None:
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start._element)
    for child in children[start_idx + 1 :]:
        body.remove(child)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if style:
        new_para.style = style
    return new_para


def add_heading_after(anchor: Paragraph, text: str, level: int) -> Paragraph:
    style = f"Heading {level}"
    return insert_paragraph_after(anchor, text, style=style)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def insert_table_after(
    doc: Document,
    anchor: Paragraph,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
) -> Paragraph:
    caption_para = insert_paragraph_after(anchor, caption, style="Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    caption_para._element.addnext(table._tbl)
    spacer = insert_paragraph_after(caption_para, "", style="Normal")
    table._tbl.addnext(spacer._element)
    return spacer


def insert_picture_group_after(anchor: Paragraph, image_path: Path, caption: str, note: str) -> Paragraph:
    image_para = insert_paragraph_after(anchor, "", style="Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.7))
    caption_para = insert_paragraph_after(image_para, caption, style="Caption")
    note_para = insert_paragraph_after(caption_para, note, style="Normal")
    return note_para


def replace_cover(paragraphs: list[Paragraph]) -> None:
    mapping = {
        "TÍTULO": "LOG MONITOR MLOPS",
        "Projeto Integrado em Engenharia Informática e Aplicações  |  LEIA  |  Data": (
            "Projeto Integrado em Engenharia Informática e Aplicações | LEIA | 28/04/2026"
        ),
        "(Nome do trabalho), Copyright de (Nome do(s) aluno(s)), IPLUSO.": (
            "LOG MONITOR MLOPS, Copyright de Raul Brito, IPLUSO."
        ),
    }
    student_slot = 0
    for paragraph in paragraphs:
        text = paragraph_text(paragraph)
        if text == "Nome do Aluno:" and student_slot == 0:
            paragraph.text = "Nome do Aluno: Raul Brito (A22309632)"
            student_slot += 1
            continue
        if text == "Nome do Aluno:" and student_slot == 1:
            paragraph.text = "Natureza do trabalho: desenvolvimento individual"
            student_slot += 1
            continue
        if text in mapping:
            paragraph.text = mapping[text]


def fill_summary_sections(doc: Document) -> None:
    resumo = find_paragraph(doc, "Resumo")
    abstract = find_paragraph(doc, "Abstract")
    indice = find_paragraph(doc, "Índice")
    lista_figuras = find_paragraph(doc, "Lista de Figuras")
    lista_tabelas = find_paragraph(doc, "Lista de Tabelas")
    problema = find_paragraph(doc, "Identificação do Problema")

    remove_between(doc, resumo, abstract)
    remove_between(doc, abstract, indice)
    remove_between(doc, lista_figuras, lista_tabelas)
    remove_between(doc, lista_tabelas, problema)

    resumo_text = (
        "O Log Monitor MLOps e um sistema de monitorizacao de logs orientado a deteccao "
        "de anomalias e ameaças em aplicacoes web. A solucao combina ingestao de logs, "
        "regras deterministicas, machine learning, tracking de experimentos e observabilidade "
        "numa stack local baseada em Docker Compose. No estado atual, o projeto ja implementa "
        "o pipeline principal de ingestao, deteccao por regras, deteccao de anomalias com "
        "Isolation Forest, dashboard de operacao, MLflow, Prometheus, Grafana, Alertmanager, "
        "CI/CD e hardening base. Permanecem parciais ou planeadas a recolha operacional de "
        "feedback via dashboard, o ciclo completo de incidentes, o benchmark externo em CICIDS-2017 "
        "e o retraining automatico. Este relatorio intermédio distingue explicitamente o que esta "
        "implementado, o que esta parcial e o que fica reservado para a fase final do projeto."
    )
    abstract_text = (
        "Log Monitor MLOps is a log-monitoring platform focused on anomaly detection and threat "
        "analysis for web applications. The current solution combines log ingestion, deterministic "
        "rules, Isolation Forest-based anomaly detection, experiment tracking, a read-only operator "
        "dashboard, and an observability stack deployed locally with Docker Compose. The implemented "
        "core already covers ingestion, rule-based alerts, MLflow tracking, monitoring, CI/CD, and "
        "baseline security hardening. The report makes a clear distinction between implemented "
        "capabilities, partially implemented features, and work planned for the final phase, namely "
        "operational feedback, full incident workflow, CICIDS-2017 benchmarking, and automatic retraining."
    )

    insert_paragraph_after(resumo, resumo_text, style="Normal")
    insert_paragraph_after(abstract, abstract_text, style="Body Text")

    figure_entries = [
        "Figura 1 - Diagrama UML dos casos de uso principais.",
    ]
    table_entries = [
        "Tabela 1 - Gap analysis do problema.",
        "Tabela 2 - Requisitos funcionais principais.",
        "Tabela 3 - Requisitos nao-funcionais principais.",
        "Tabela 4 - Resumo de use cases e estado de implementacao.",
        "Tabela 5 - Estado atual do trabalho face ao planeamento.",
        "Tabela 6 - Comparacao de benchmarking.",
        "Tabela 7 - Fases, atividades e proximos passos.",
        "Tabela 8 - Gantt resumido por blocos de semanas.",
        "Tabela 9 - Atividades realizadas e a realizar.",
    ]
    current = lista_figuras
    for entry in figure_entries:
        current = insert_paragraph_after(current, entry, style="Normal")
    current = lista_tabelas
    for entry in table_entries:
        current = insert_paragraph_after(current, entry, style="Normal")


def fill_problem_section(doc: Document) -> None:
    start = find_paragraph(doc, "Identificação do Problema")
    end = find_paragraph(doc, "Levantamento e análise de requisitos")
    remove_between(doc, start, end)

    current = insert_paragraph_after(
        start,
        (
            "O projeto parte de um problema real: equipas pequenas e contextos academicos "
            "ou laboratoriais raramente conseguem operar ferramentas SIEM comerciais devido ao "
            "custo, complexidade e baixa adaptabilidade a cenarios experimentais. Ao mesmo tempo, "
            "logs HTTP continuam a ser uma fonte critica para detetar brute force, scanning, "
            "SQL injection, path traversal e anomalias comportamentais."
        ),
        style="Normal",
    )
    current = insert_paragraph_after(
        current,
        (
            "O objetivo do Log Monitor MLOps e disponibilizar uma stack local, reprodutivel e "
            "explicavel, capaz de gerar, ingerir, armazenar, analisar e visualizar logs em tempo "
            "quase real. O foco da fase intermédia nao e apresentar um produto final acabado, mas "
            "demonstrar um nucleo tecnico consistente, validavel e com espaco claro para evolucao."
        ),
        style="Normal",
    )
    insert_table_after(
        doc,
        current,
        "Tabela 1 - Gap analysis do problema.",
        ["Lacuna", "Impacto", "Resposta no projeto"],
        [
            ["Explicabilidade fraca em ML", "Reduz confianca operacional", "SHAP offline ja produzido; explicacao operacional continua parcial"],
            ["Feedback humano pouco estruturado", "Dificulta melhoria continua", "Schema de feedback preparado; workflow completo fica para fase final"],
            ["Custo e lock-in de SIEM comercial", "Barreira para PMEs e ensino", "Stack local open-source baseada em Docker Compose"],
        ],
    )


def fill_requirements_section(doc: Document) -> None:
    start = find_paragraph(doc, "Levantamento e análise de requisitos")
    end = find_paragraph(doc, "Viabilidade e Pertinência")
    remove_between(doc, start, end)

    current = insert_paragraph_after(
        start,
        "Os requisitos abaixo resumem o scope intermédio e identificam explicitamente o estado de cada capacidade.",
        style="Normal",
    )
    current = add_heading_after(current, "2.1 Requisitos Funcionais", 2)
    current = insert_table_after(
        doc,
        current,
        "Tabela 2 - Requisitos funcionais principais.",
        ["ID", "Requisito", "Estado"],
        [
            ["RF-01", "Ingerir logs HTTP e persisti-los em PostgreSQL/TimescaleDB", "Implementado"],
            ["RF-02", "Detetar ataques conhecidos por regras deterministicas", "Implementado"],
            ["RF-03", "Detetar anomalias desconhecidas com Isolation Forest", "Implementado"],
            ["RF-04", "Combinar sinais de regras e ML num pipeline hibrido", "Parcial"],
            ["RF-05", "Expor resultados em dashboard de operador", "Implementado"],
            ["RF-06", "Registar feedback humano e suportar retraining", "Parcial / Planeado"],
        ],
    )
    current = add_heading_after(current, "2.2 Requisitos Nao-Funcionais", 2)
    current = insert_table_after(
        doc,
        current,
        "Tabela 3 - Requisitos nao-funcionais principais.",
        ["ID", "Requisito", "Evidencia atual"],
        [
            ["RNF-01", "Performance de ingestao", "10.933 logs/s com batch 500"],
            ["RNF-02", "Observabilidade e alerting", "Prometheus, Grafana e Alertmanager implementados"],
            ["RNF-03", "Reprodutibilidade", "Docker Compose + MLflow"],
            ["RNF-04", "Qualidade e testes", "87/102 testes a passar; 70.97% coverage"],
            ["RNF-05", "Seguranca base", "Validacao, headers, CI de seguranca e hardening inicial"],
        ],
    )
    current = add_heading_after(current, "2.3 Atores do Sistema", 2)
    current = insert_paragraph_after(
        current,
        (
            "Os atores principais sao quatro: Security Analyst, responsavel pela investigacao "
            "de alertas; Developer, focado na evolucao tecnica da plataforma; Administrator, "
            "encarregado da operacao e configuracao; e o Sistema (Automated), que gera logs, "
            "corre regras, calcula scores e publica metricas."
        ),
        style="Normal",
    )
    current = add_heading_after(current, "2.4 Casos de Uso UML", 2)
    current = insert_picture_group_after(
        current,
        combine_images(),
        "Figura 1 - Diagrama UML dos casos de uso principais.",
        (
            "O diagrama inclui os casos de uso essenciais das categorias P0, P1 e P2. "
            "Na fase intermédia, os casos de uso ligados a ingestao, regras, treino, pipeline "
            "ML, observabilidade e dashboard estao concluídos ou demonstraveis; feedback "
            "operacional completo e benchmark CICIDS permanecem fora do scope fechado."
        ),
    )
    current = add_heading_after(current, "2.5 Resumo de Use Cases por Prioridade", 2)
    insert_table_after(
        doc,
        current,
        "Tabela 4 - Resumo de use cases e estado de implementacao.",
        ["Use case", "Prioridade", "Estado"],
        [
            ["Gerar logs sinteticos e ingerir logs", "P0", "Implementado"],
            ["Detetar ataque por regras", "P0", "Implementado"],
            ["Treinar modelos ML", "P0", "Implementado"],
            ["Executar pipeline hibrido", "P0", "Parcial"],
            ["Visualizar dashboard", "P0", "Implementado"],
            ["Marcar falso positivo", "P1", "Parcial / Planeado"],
            ["Gerir incidente", "P1", "Parcial"],
            ["Validar em CICIDS", "P2", "Planeado"],
        ],
    )


def fill_viability_section(doc: Document) -> None:
    start = find_paragraph(doc, "Viabilidade e Pertinência")
    end = find_paragraph(doc, "Solução Proposta")
    remove_between(doc, start, end)

    current = add_heading_after(start, "3.1 Viabilidade tecnica e economica", 2)
    current = insert_paragraph_after(
        current,
        (
            "A viabilidade tecnica ja foi demonstrada pela integracao funcional entre Flask, "
            "PostgreSQL/TimescaleDB, pipeline de regras, pipeline ML, MLflow, dashboard Streamlit "
            "e stack de monitorizacao. O sistema corre localmente por Docker Compose, o que reduz "
            "friccao de instalacao e facilita repetibilidade."
        ),
        style="Normal",
    )
    current = insert_paragraph_after(
        current,
        (
            "A viabilidade economica decorre da opcao por software open-source e infraestrutura "
            "local. O custo principal passa a ser tempo de desenvolvimento e operacao, o que e "
            "adequado ao contexto academico e a cenarios de pequena escala onde um SIEM comercial "
            "seria desproporcionado."
        ),
        style="Normal",
    )
    current = add_heading_after(current, "3.2 Pertinencia academica e pratica", 2)
    current = insert_paragraph_after(
        current,
        (
            "O projeto e pertinente porque liga areas que normalmente aparecem separadas: "
            "engenharia de software, bases de dados, MLOps, monitorizacao e ciberseguranca. "
            "Na pratica, cria uma plataforma demonstravel para analisar logs, testar regras, "
            "medir desempenho e estudar compromissos entre deteccao supervisionada e nao supervisionada."
        ),
        style="Normal",
    )
    current = add_heading_after(current, "3.3 Continuidade apos a fase intermédia", 2)
    insert_paragraph_after(
        current,
        (
            "O trabalho nao se esgota nesta entrega. Permanecem claramente identificados como "
            "proximos passos o fecho do feedback operacional, o benchmark externo, o retraining "
            "assistido por feedback e o refinamento do workflow de incidentes. Esta separacao "
            "entre realizado, parcial e planeado torna a continuidade do projeto objetiva e gerivel."
        ),
        style="Normal",
    )


def fill_solution_section(doc: Document) -> None:
    start = find_paragraph(doc, "Solução Proposta")
    end = find_paragraph(doc, "Benchmarking")
    remove_between(doc, start, end)

    current = add_heading_after(start, "4.1 Estado atual do trabalho face ao planeamento", 2)
    current = insert_table_after(
        doc,
        current,
        "Tabela 5 - Estado atual do trabalho face ao planeamento.",
        ["Area", "Estado", "Evidencia atual"],
        [
            ["Geracao, ingestao e armazenamento", "Implementado", "Flask, ingester e PostgreSQL/TimescaleDB"],
            ["Deteccao por regras", "Implementado", "6 regras ativas no rule engine"],
            ["ML e tracking", "Implementado", "Isolation Forest, RandomForest e MLflow"],
            ["Dashboard de operador", "Implementado", "Streamlit read-only com overview, alerts e log explorer"],
            ["Feedback humano", "Parcial", "Schema preparado; fluxo operacional ainda nao fechado"],
            ["Incident workflow", "Parcial", "Playbooks estaticos; sem ciclo NEW->RESOLVED completo"],
            ["Benchmark CICIDS-2017", "Planeado", "Arquitetura preparada; integracao ainda nao feita"],
            ["Retraining automatico", "Planeado", "Dependente de feedback validado"],
        ],
    )
    current = add_heading_after(current, "4.2 Arquitetura e componentes", 2)
    current = insert_paragraph_after(
        current,
        (
            "A arquitetura atual organiza-se em quatro camadas: geracao e ingestao de logs; "
            "deteccao e scoring; armazenamento e tracking; observabilidade e operacao. "
            "Na pratica, a stack integra Flask, PostgreSQL/TimescaleDB, rule engine, pipeline ML, "
            "MLflow, Streamlit, Prometheus, Grafana e Alertmanager."
        ),
        style="Normal",
    )
    current = insert_paragraph_after(
        current,
        (
            "Do ponto de vista funcional, o fluxo principal ja demonstravel e o seguinte: a "
            "aplicacao Flask gera pedidos HTTP sinteticos; o ingester persiste os eventos; o "
            "rule engine verifica ataques conhecidos; o pipeline ML calcula scores de anomalia; "
            "MLflow regista experiencias e artefactos; e o dashboard apresenta logs, alertas e "
            "metricas operacionais. O valor desta composicao esta em reunir num unico projeto "
            "componentes que normalmente seriam analisados isoladamente."
        ),
        style="Normal",
    )
    current = add_heading_after(current, "4.3 Resultados mensurados e classificacao funcional", 2)
    current = insert_paragraph_after(
        current,
        (
            "Os resultados mais relevantes ja medidos foram: 10.933 logs/s de ingestao com "
            "batch 500; pipeline ML treinado em menos de um minuto no dataset interno; score "
            "operacional do modelo hibrido acima da meta de 0.75 no holdout interno; e stack "
            "de observabilidade a recolher metricas de logs, alertas e F1. Em QA, a execucao "
            "local mais recente do pytest recolheu 102 testes, com 87 a passar, 1 skipped, "
            "2 falhas, 12 erros e cobertura global de 70.97%, concentrando os problemas atuais "
            "no modulo hybrid_pipeline."
        ),
        style="Normal",
    )
    current = insert_paragraph_after(
        current,
        (
            "Assim, para esta entrega, considera-se implementado tudo o que pode ser demonstrado "
            "de ponta a ponta em ambiente local; parcial o que existe em estrutura ou MVP sem "
            "workflow fechado; e planeado o que foi deliberadamente adiado para a fase final."
        ),
        style="Normal",
    )
    insert_paragraph_after(
        current,
        (
            "Esta leitura e importante para a apresentacao intermédia: o projeto ja tem um nucleo "
            "operacional forte para demonstracao, mas ainda nao deve ser descrito como plataforma "
            "totalmente fechada em feedback humano, ciclo de incidente ou benchmark externo."
        ),
        style="Normal",
    )


def fill_benchmarking_section(doc: Document) -> None:
    start = find_paragraph(doc, "Benchmarking")
    end = find_paragraph(doc, "Método e planeamento")
    remove_between(doc, start, end)

    current = insert_paragraph_after(
        start,
        (
            "A comparacao com alternativas existentes ajuda a enquadrar o valor do projeto. "
            "O objetivo nao e competir em escala com suites empresariais, mas demonstrar "
            "um equilibrio entre custo, reprodutibilidade, observabilidade e explicabilidade."
        ),
        style="Normal",
    )
    current = insert_table_after(
        doc,
        current,
        "Tabela 6 - Comparacao de benchmarking.",
        ["Solucao", "Custo", "Explainability", "Feedback", "Adequacao ao projeto"],
        [
            ["Splunk ES", "Muito alto", "Baixa a media", "Proprietario", "Baixa para contexto academico"],
            ["Datadog Security", "Alto", "Media", "Limitado", "Boa UX, mas forte dependência SaaS"],
            ["Elastic SIEM", "Medio", "Baixa", "Inexistente", "Flexivel, mas mais pesado de operar"],
            ["Log Monitor MLOps", "Baixo", "SHAP offline; explicacao operacional parcial", "Parcial", "Alta para ensino, demo e experimentacao"],
        ],
    )
    current = insert_paragraph_after(
        current,
        (
            "Face ao mercado, a principal mais-valia do projeto esta na integracao de stack "
            "local, tracking de ML, regras, dashboard e monitorizacao numa solucao unica e "
            "explicavel, ainda que com menos maturidade operacional do que plataformas comerciais."
        ),
        style="Normal",
    )
    insert_paragraph_after(
        current,
        (
            "Em termos de posicionamento, o projeto ganha sobretudo em controlo do ambiente, "
            "baixo custo e transparencia tecnica. Perde, naturalmente, em funcionalidades empresariais "
            "como conectores multiplos maduros, automacao de incident response e validacao extensa "
            "em datasets externos. Esta diferenca nao fragiliza o projeto; pelo contrario, ajuda a "
            "definir com rigor o que esta a ser demonstrado nesta fase."
        ),
        style="Normal",
    )


def fill_planning_section(doc: Document) -> None:
    start = find_paragraph(doc, "Método e planeamento")
    end = find_paragraph(doc, "Bibliografia")
    remove_between(doc, start, end)

    current = add_heading_after(start, "6.1 Abordagem de desenvolvimento", 2)
    current = insert_paragraph_after(
        current,
        (
            "O projeto seguiu uma abordagem incremental por fases, com validacao progressiva "
            "do nucleo tecnico antes de expandir para observabilidade, dashboard, seguranca e CI/CD."
        ),
        style="Normal",
    )
    current = insert_table_after(
        doc,
        current,
        "Tabela 7 - Fases, atividades e proximos passos.",
        ["Fase", "Periodo", "Objetivo principal", "Estado"],
        [
            ["Goal A", "Semanas 1-8", "MVP de ingestao, regras, ML e explainability base", "Concluido"],
            ["Goal B", "Semanas 9-14", "Docker, monitorizacao, dashboard, testes e seguranca base", "Maioritariamente concluido"],
            ["Goal C", "Semanas 15-16", "Consolidacao da demo, documentacao e estabilizacao", "Em curso"],
            ["Fase final", "Apos intercalar", "Feedback operacional, benchmark externo e retraining", "Planeado"],
        ],
    )
    current = add_heading_after(current, "6.2 Timeline e Gantt resumido", 2)
    current = insert_paragraph_after(
        current,
        (
            "A tabela seguinte recompõe o cronograma numa vista Gantt simplificada, suficiente "
            "para mostrar a distribuicao temporal das frentes principais sem ocupar o espaco de "
            "uma matriz semanal extensa."
        ),
        style="Normal",
    )
    current = insert_table_after(
        doc,
        current,
        "Tabela 8 - Gantt resumido por blocos de semanas.",
        ["Frente", "S1-2", "S3-4", "S5-6", "S7-8", "S9-10", "S11-12", "S13-14", "S15-16"],
        [
            ["Setup e arquitetura base", "X", "X", "", "", "", "", "", ""],
            ["Ingestao e persistencia", "", "X", "X", "", "", "", "", ""],
            ["Regras e ML core", "", "", "X", "X", "", "", "", ""],
            ["Dashboard e observabilidade", "", "", "", "", "X", "X", "", ""],
            ["QA, CI/CD e hardening", "", "", "", "", "", "X", "X", ""],
            ["Documentacao e demo final", "", "", "", "", "", "", "", "X"],
        ],
    )
    current = add_heading_after(current, "6.3 Atividades realizadas e a realizar", 2)
    current = insert_table_after(
        doc,
        current,
        "Tabela 9 - Atividades realizadas e a realizar.",
        ["Realizadas ate agora", "A realizar na fase final"],
        [
            ["Geracao e ingestao de logs", "Fechar feedback via dashboard"],
            ["6 regras de deteccao", "Melhorar QA do hybrid_pipeline"],
            ["Treino ML com MLflow", "Executar benchmark CICIDS-2017 ou equivalente"],
            ["Dashboard Streamlit e observabilidade", "Refinar lifecycle de incidentes"],
            ["CI/CD e hardening base", "Consolidar documentacao final e demo"],
        ],
    )
    current = add_heading_after(current, "6.4 Responsabilidades do autor", 2)
    current = insert_paragraph_after(
        current,
        (
            "Tratando-se de um trabalho individual, o desenvolvimento foi assegurado pelo autor "
            "em quatro frentes principais: engenharia de dados e persistencia; deteccao e ML; "
            "observabilidade e dashboard; e qualidade, seguranca e documentacao. Esta divisao "
            "por areas permite justificar o planeamento mesmo sem distribuicao por varios membros."
        ),
        style="Normal",
    )
    current = add_heading_after(current, "6.5 Riscos e mitigacao", 2)
    insert_paragraph_after(
        current,
        (
            "Os riscos principais identificados nesta fase sao tres: desalinhamento entre o "
            "relatorio e o estado real do codigo; instabilidade residual no modulo hybrid_pipeline; "
            "e expansao excessiva de scope antes da entrega final. A mitigacao definida passa por "
            "manter classificacao implementado/parcial/planeado em todas as secoes, priorizar a "
            "correcao dos testes mais criticos e diferir funcionalidades de baixo ROI."
        ),
        style="Normal",
    )


def fill_bibliography(doc: Document) -> None:
    start = find_paragraph(doc, "Bibliografia")
    remove_after(doc, start)
    refs = [
        "[1] Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation Forest. IEEE International Conference on Data Mining.",
        "[2] Lundberg, S. M., & Lee, S. I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS 30.",
        "[3] Zaharia, M., Chen, A., Davidson, A., et al. (2018). Accelerating the Machine Learning Lifecycle with MLflow. IEEE Data Engineering Bulletin, 41(4).",
        "[4] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12.",
        "[5] Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization.",
        "[6] Docker Inc. (2026). Docker Compose documentation. https://docs.docker.com/",
        "[7] Prometheus Authors. (2026). Prometheus documentation. https://prometheus.io/docs/",
        "[8] Grafana Labs. (2026). Grafana documentation. https://grafana.com/docs/",
    ]
    current = start
    for ref in refs:
        current = insert_paragraph_after(current, ref, style="Normal")


def main() -> None:
    TEMPLATE.replace(OUTPUT) if False else None
    doc = Document(TEMPLATE)
    replace_cover(doc.paragraphs)
    fill_summary_sections(doc)
    fill_problem_section(doc)
    fill_requirements_section(doc)
    fill_viability_section(doc)
    fill_solution_section(doc)
    fill_benchmarking_section(doc)
    fill_planning_section(doc)
    fill_bibliography(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
